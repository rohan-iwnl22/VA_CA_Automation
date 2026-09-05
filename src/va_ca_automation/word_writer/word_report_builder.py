"""Build the Word audit report from template and processed data."""

from __future__ import annotations

import io
import logging
import os
import shutil
import tempfile
from copy import deepcopy
from math import cos, radians, sin
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from ..metadata.engagement_metadata import EngagementMetadata

logger = logging.getLogger("va_ca_automation")

VA_COLUMNS = [
    "Sr. no",
    "Vulnerbility Title",
    "Description",
    "Risk",
    "Host",
    "Port",
    "Recommendation ",
    "Reference",
    "CVE",
]
CA_COLUMNS = ["Sr.No.", "Title", "Host", "Description", "Solution", "Risk"]

RISK_COLORS = {
    "Critical": "#C00000",
    "High": "#FF0000",
    "Medium": "#FFC000",
    "Low": "#FFFF00",
    "Info": "#00B0F0",
    "Failed": "#4472C4",
    "Warning": "#ED7D31",
}

CA_RISK_COLORS = {
    "FAILED": "#4472C4",
    "WARNING": "#ED7D31",
}


def _set_cell_shading(cell, color_hex: str) -> None:
    """Set background shading on a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _set_row_height(row, height_pt: float) -> None:
    """Set row height in points."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(height_pt * 20)))
    trHeight.set(qn("w:hRule"), "atLeast")
    trPr.append(trHeight)


def _repeat_table_header(row) -> None:
    """Mark a table row as a repeating header row for page breaks."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def _prevent_row_split(row) -> None:
    """Keep a row together when Word paginates the document."""
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def _set_cell_borders(cell, color: str = "000000", size: str = "4") -> None:
    """Set thin borders on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)
        tcBorders.append(element)
    tcPr.append(tcBorders)


def _set_cell_vertical_alignment(cell, align: str = "center") -> None:
    """Set vertical alignment on a table cell (top, center, bottom)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), align)
    tcPr.append(vAlign)


def _set_table_width(table, width_cm: float) -> None:
    """Set the table width in centimeters."""
    # 1 cm = 567 DXA (twentieths of a point)
    width_dxa = int(width_cm * 567)
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(width_dxa))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def _set_table_fixed_layout(table) -> None:
    """Set table layout to fixed so column widths are respected."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def _set_column_widths(table, widths_dxa: list[int]) -> None:
    """Set individual column widths in DXA units for the grid and all cells."""
    grid_cols = table._tbl.tblGrid.gridCol_lst
    for grid_col, width in zip(grid_cols, widths_dxa):
        grid_col.set(qn("w:w"), str(width))

    ns = qn("w:tcW")
    for row in table.rows:
        for j, cell in enumerate(row.cells):
            if j < len(widths_dxa):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = tcPr.find(ns)
                if tcW is None:
                    tcW = OxmlElement("w:tcW")
                    tcPr.append(tcW)
                tcW.set(qn("w:w"), str(widths_dxa[j]))
                tcW.set(qn("w:type"), "dxa")


def _usable_page_width_cm(doc: Document) -> float:
    """Return the printable width of the report's active page layout."""
    section = doc.sections[-1]
    return (section.page_width - section.left_margin - section.right_margin) / Cm(1)


def _configure_table(doc: Document, table, proportions: list[int]) -> None:
    """Fit a fixed-layout table to the template's printable landscape width."""
    usable_width_cm = _usable_page_width_cm(doc)
    # Keep a small gutter so Word does not push the right edge into the margin.
    table_width_cm = max(1.0, usable_width_cm - 0.2)
    total_dxa = int(table_width_cm * 567)
    widths_dxa = [round(total_dxa * value / sum(proportions)) for value in proportions]
    widths_dxa[-1] += total_dxa - sum(widths_dxa)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_width(table, table_width_cm)
    _set_table_fixed_layout(table)
    _set_column_widths(table, widths_dxa)


def _generate_pie_chart_image(
    risk_summary: dict[str, int],
    title: str,
    colors: dict[str, str],
    ca_risk_summary: dict[str, int] | None = None,
) -> io.BytesIO:
    """Generate a pie chart image and return as BytesIO.

    If ca_risk_summary is provided, FAILED and WARNING counts are included
    in the pie chart alongside the VA risk levels.
    """
    labels = []
    sizes = []
    chart_colors = []

    # Add VA risk levels (Critical, High, Medium, Low, Info)
    for label, count in risk_summary.items():
        if label == "Grand Total" or count == 0:
            continue
        labels.append(label)
        sizes.append(count)
        chart_colors.append(colors.get(label, "#999999"))

    # Add CA risk levels (FAILED, WARNING) if provided
    if ca_risk_summary:
        for label, count in ca_risk_summary.items():
            if label == "Grand Total" or count == 0:
                continue
            # Map FAILED -> Failed, WARNING -> Warning for display
            display_label = label.title() if label.isupper() else label
            labels.append(display_label)
            sizes.append(count)
            chart_colors.append(colors.get(display_label, colors.get(label, "#999999")))

    if not sizes:
        return None

    fig, ax = plt.subplots(figsize=(5, 4))
    wedges, _ = ax.pie(
        sizes,
        labels=None,
        colors=chart_colors,
        startangle=90,
    )
    total = sum(sizes)
    for wedge, value in zip(wedges, sizes):
        angle = radians((wedge.theta1 + wedge.theta2) / 2)
        # Values on narrow slices overlap when placed inside the pie. Put those
        # values just outside the slice, retaining number-only chart content.
        radius = 0.6 if value / total >= 0.10 else 1.23
        x, y = radius * cos(angle), radius * sin(angle)
        ax.text(
            x,
            y,
            str(value),
            ha="center",
            va="center",
            fontsize=10 if radius < 1 else 9,
            fontweight="bold",
        )
    ax.set_xlim(-1.4, 1.4)
    ax.set_ylim(-1.4, 1.4)
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _replace_text_in_paragraph(para, old_text: str, new_text: str) -> None:
    """Replace text in a paragraph, handling text split across multiple runs."""
    full_text = para.text
    if old_text not in full_text:
        return

    # Strategy: clear all runs, set the replaced text on the first run
    new_full_text = full_text.replace(old_text, new_text)
    if para.runs:
        para.runs[0].text = new_full_text
        for run in para.runs[1:]:
            run.text = ""


def _replace_client_name(doc: Document, client_name: str) -> None:
    """Replace 'Client name' with the actual client name throughout the document."""
    # Replace in all paragraphs (handles text split across runs)
    for para in doc.paragraphs:
        if "Client name" in para.text or "client name" in para.text:
            _replace_text_in_paragraph(para, "Client name", client_name)
            _replace_text_in_paragraph(para, "client name", client_name)

    # Replace in all table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if "Client name" in para.text or "client name" in para.text:
                        _replace_text_in_paragraph(para, "Client name", client_name)
                        _replace_text_in_paragraph(para, "client name", client_name)

    # Replace client name in footers (including tables in footers)
    for section in doc.sections:
        footer = section.footer
        # Access footer XML directly to find all text elements
        from lxml import etree
        nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        for t_elem in footer._element.findall(".//w:t", nsmap):
            if t_elem.text and ("Client name" in t_elem.text or "client name" in t_elem.text):
                t_elem.text = t_elem.text.replace("Client name", client_name)
                t_elem.text = t_elem.text.replace("client name", client_name)


def _populate_document_details(doc: Document, metadata: EngagementMetadata) -> None:
    """Populate the Document Preparation table (Table 2) with metadata values."""
    replacements = {
        "Released Date": metadata.released_date,
        "Release date": metadata.released_date,
    }
    for table in doc.tables:
        headers = [cell.text.strip() for cell in table.rows[0].cells] if table.rows else []
        is_doc_prep = any("Document Preparation" in h for h in headers)
        for row in table.rows:
            for ci, cell in enumerate(row.cells):
                if is_doc_prep and ci == 0:
                    continue
                for para in cell.paragraphs:
                    for placeholder, value in replacements.items():
                        if placeholder in para.text and value:
                            _replace_text_in_paragraph(para, placeholder, value)
    for para in doc.paragraphs:
        for placeholder, value in replacements.items():
            if placeholder in para.text and value:
                _replace_text_in_paragraph(para, placeholder, value)

    # Table 0: Report Release Date
    if len(doc.tables) > 0:
        tbl0 = doc.tables[0]
        # Row 0, Column 1: Report Release Date value
        if len(tbl0.rows) > 0 and metadata.report_release_date:
            tbl0.rows[0].cells[1].text = metadata.report_release_date
        # Row 3, Column 1: Period
        if len(tbl0.rows) > 3 and metadata.period:
            tbl0.rows[3].cells[1].text = metadata.period

    # Table 2: Document Preparation
    if len(doc.tables) > 2:
        tbl2 = doc.tables[2]
        # Row 2: Document ID
        if len(tbl2.rows) > 2 and metadata.document_id:
            tbl2.rows[2].cells[1].text = metadata.document_id
        # Row 3: Document Version
        if len(tbl2.rows) > 3 and metadata.document_version:
            tbl2.rows[3].cells[1].text = metadata.document_version
        # Row 4: Prepared by (from Excel security_tester)
        if len(tbl2.rows) > 4 and metadata.security_tester:
            tbl2.rows[4].cells[1].text = metadata.security_tester
        # Row 5: Reviewed by
        if len(tbl2.rows) > 5 and metadata.reviewed_by:
            tbl2.rows[5].cells[1].text = metadata.reviewed_by
        # Row 6: Approved by
        if len(tbl2.rows) > 6 and metadata.approved_by:
            tbl2.rows[6].cells[1].text = metadata.approved_by
        # Row 7: Released by (use reviewed_by)
        if len(tbl2.rows) > 7 and metadata.reviewed_by:
            tbl2.rows[7].cells[1].text = metadata.reviewed_by
        # Row 12: Release date
        if len(tbl2.rows) > 12 and metadata.released_date:
            tbl2.rows[12].cells[1].text = metadata.released_date

    # Table 3: Document Change History (Row 2 = data row)
    if len(doc.tables) > 3:
        tbl3 = doc.tables[3]
        if len(tbl3.rows) > 2:
            if metadata.change_history_version:
                tbl3.rows[2].cells[0].text = metadata.change_history_version
            if metadata.change_history_date:
                tbl3.rows[2].cells[1].text = metadata.change_history_date
            if metadata.change_history_remarks:
                tbl3.rows[2].cells[2].text = metadata.change_history_remarks
            # Center align all cells in Row 2
            for cell in tbl3.rows[2].cells:
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.size = Pt(12)
                        run.font.name = "Cambria"

    # Table 4: Document Distribution List (Row 2 = data row)
    if len(doc.tables) > 4:
        tbl4 = doc.tables[4]
        if len(tbl4.rows) > 2:
            if metadata.distribution_name:
                tbl4.rows[2].cells[0].text = metadata.distribution_name
            if metadata.distribution_organization:
                tbl4.rows[2].cells[1].text = metadata.distribution_organization
            if metadata.distribution_designation:
                tbl4.rows[2].cells[2].text = metadata.distribution_designation
            if metadata.distribution_email:
                tbl4.rows[2].cells[3].text = metadata.distribution_email
            # Center align all cells in Row 2
            for cell in tbl4.rows[2].cells:
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.size = Pt(12)
                        run.font.name = "Cambria"

    # Table 8: Details of Auditing Team (Row 1 & 2: user input, Row 3: default)
    if len(doc.tables) > 8:
        tbl8 = doc.tables[8]
        # Row 1: First auditor (user input)
        if len(tbl8.rows) > 1:
            if metadata.auditor_1_name:
                tbl8.rows[1].cells[1].text = metadata.auditor_1_name
            if metadata.auditor_1_designation:
                tbl8.rows[1].cells[2].text = metadata.auditor_1_designation
            if metadata.auditor_1_email:
                tbl8.rows[1].cells[3].text = metadata.auditor_1_email
            if metadata.auditor_1_qualifications:
                tbl8.rows[1].cells[4].text = metadata.auditor_1_qualifications
            if metadata.auditor_1_cert_in:
                tbl8.rows[1].cells[5].text = metadata.auditor_1_cert_in
        # Row 2: Second auditor (user input)
        if len(tbl8.rows) > 2:
            if metadata.auditor_2_name:
                tbl8.rows[2].cells[1].text = metadata.auditor_2_name
            if metadata.auditor_2_designation:
                tbl8.rows[2].cells[2].text = metadata.auditor_2_designation
            if metadata.auditor_2_email:
                tbl8.rows[2].cells[3].text = metadata.auditor_2_email
            if metadata.auditor_2_qualifications:
                tbl8.rows[2].cells[4].text = metadata.auditor_2_qualifications
            if metadata.auditor_2_cert_in:
                tbl8.rows[2].cells[5].text = metadata.auditor_2_cert_in
        # Row 3: Default (Mr. Rohan Sawant)
        if len(tbl8.rows) > 3:
            tbl8.rows[3].cells[1].text = "Mr. Rohan Sawant"
            tbl8.rows[3].cells[2].text = "Senior Manager - IT Security"
            tbl8.rows[3].cells[3].text = "rohan.sawant@secunatix.com"
            tbl8.rows[3].cells[4].text = "CEH"
            tbl8.rows[3].cells[5].text = "Yes"


def _populate_executive_summary_table(doc: Document, va_risk: dict[str, int], ca_risk: dict[str, int]) -> None:
    """Populate Table 11 (Security Assessment summary) and Table 13 (Risk Classification)."""
    # Table 11: Security Assessment summary (5 rows x 9 cols)
    # Row 0: merged header "Security Assessment" / "Initial VAPT Scan"
    # Row 1: headers Sr. No., VAPT, Critical, High, Medium, Low, Info, Failed, Warning
    # Row 2: VA data
    # Row 3: CA data
    # Row 4: Total
    if len(doc.tables) > 11:
        tbl = doc.tables[11]
        # Row 2 (VA): fill Critical, High, Medium, Low
        va_row = tbl.rows[2]
        va_row.cells[2].text = str(va_risk.get("Critical", 0))
        va_row.cells[3].text = str(va_risk.get("High", 0))
        va_row.cells[4].text = str(va_risk.get("Medium", 0))
        va_row.cells[5].text = str(va_risk.get("Low", 0))
        va_row.cells[6].text = "0"
        va_row.cells[7].text = "-"
        va_row.cells[8].text = "-"

        # Row 3 (CA): fill Failed, Warning
        ca_row = tbl.rows[3]
        ca_row.cells[2].text = "-"
        ca_row.cells[3].text = "-"
        ca_row.cells[4].text = "-"
        ca_row.cells[5].text = "-"
        ca_row.cells[6].text = "0"
        ca_row.cells[7].text = str(ca_risk.get("FAILED", 0))
        ca_row.cells[8].text = str(ca_risk.get("WARNING", 0))

        # Row 4 (Total)
        total_row = tbl.rows[4]
        total_row.cells[2].text = str(va_risk.get("Critical", 0))
        total_row.cells[3].text = str(va_risk.get("High", 0))
        total_row.cells[4].text = str(va_risk.get("Medium", 0))
        total_row.cells[5].text = str(va_risk.get("Low", 0))
        total_row.cells[6].text = "0"
        total_row.cells[7].text = str(ca_risk.get("FAILED", 0))
        total_row.cells[8].text = str(ca_risk.get("WARNING", 0))

    # Table 13: Risk Classification (8 rows x 2 cols)
    if len(doc.tables) > 13:
        tbl13 = doc.tables[13]
        risk_map = {
            0: "Critical",
            1: "High",
            2: "Medium",
            3: "Low",
            4: "Info",
            5: "Failed",
            6: "Compliance",
            7: "Total",
        }
        for idx, label in risk_map.items():
            if label in ("Critical", "High", "Medium", "Low", "Info"):
                val = va_risk.get(label, 0)
            elif label == "Failed":
                val = ca_risk.get("FAILED", 0)
            elif label == "Compliance":
                val = ca_risk.get("WARNING", 0)
            elif label == "Total":
                val = sum(v for k, v in va_risk.items() if k != "Grand Total") + \
                      sum(v for k, v in ca_risk.items() if k != "Grand Total")
            else:
                val = 0
            tbl13.rows[idx].cells[1].text = str(val)


def _populate_executive_summary_text(doc: Document, va_risk: dict[str, int], ca_risk: dict[str, int]) -> None:
    """Update the executive summary paragraph with actual counts."""
    total_ca_failed = ca_risk.get("FAILED", 0)
    total_ca_warning = ca_risk.get("WARNING", 0)

    for para in doc.paragraphs:
        if "The audit reported" in para.text and "Critical" in para.text:
            new_text = (
                f"The audit reported {va_risk.get('Critical', 0)} Critical, "
                f"{va_risk.get('High', 0)} High, {va_risk.get('Medium', 0)} Medium, "
                f"{va_risk.get('Low', 0)} Low risk, 0 informational vulnerabilities and "
                f"{total_ca_failed} Failed, {total_ca_warning} Warning compliance issue "
                f"on the target asset. Following table and graph summarizes overall risk "
                f"of target infrastructure."
            )
            # Clear all runs and set text on first run
            if para.runs:
                para.runs[0].text = new_text
                for run in para.runs[1:]:
                    run.text = ""
            break


def _value_for_column(row: pd.Series, column: str) -> object:
    """Read a report value while accepting the source workbooks' header variants."""
    aliases = {
        "Vulnerbility Title": ("Vulnerbility Title", "Vulnerability Title"),
        "Sr.No.": ("Sr.No.", "Sr. no", "Sr. No."),
        "Solution": ("Solution", "Solution "),
        "Risk": ("Risk", "Status"),
    }
    for source_column in aliases.get(column, (column,)):
        if source_column in row.index:
            return row[source_column]
    return ""


def _create_detailed_table(
    doc: Document,
    anchor_paragraph,
    data: pd.DataFrame,
    columns: list[str],
    proportions: list[int],
    narrative_columns: set[str],
    top_left_columns: set[str] | None = None,
) -> None:
    """Insert one report table at its textual anchor using the template page geometry."""
    if top_left_columns is None:
        top_left_columns = set()

    table = doc.add_table(rows=len(data) + 1, cols=len(columns))
    anchor_paragraph._p.addnext(table._tbl)
    _configure_table(doc, table, proportions)

    header_row = table.rows[0]
    _repeat_table_header(header_row)
    _set_row_height(header_row, 24)
    for index, column in enumerate(columns):
        cell = header_row.cells[index]
        cell.text = column
        _set_cell_shading(cell, "FFC000")
        _set_cell_borders(cell)
        _set_cell_vertical_alignment(cell)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(12)
                run.font.name = "Cambria"

    for row_index, (_, row) in enumerate(data.iterrows(), start=1):
        data_row = table.rows[row_index]
        _set_row_height(data_row, 18)
        for column_index, column in enumerate(columns):
            cell = data_row.cells[column_index]
            value = _value_for_column(row, column)
            cell.text = "N/A" if pd.isna(value) or value == "" else str(value)
            _set_cell_borders(cell)

            if column in top_left_columns:
                _set_cell_vertical_alignment(cell, "top")
            else:
                _set_cell_vertical_alignment(cell)

            for paragraph in cell.paragraphs:
                if column in top_left_columns:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif column in narrative_columns:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(12)
                    run.font.name = "Cambria"

            if column == "Risk":
                risk = str(value).strip().upper() if not pd.isna(value) else ""
                is_ca_table = "Sr.No." in columns
                if is_ca_table:
                    if risk in CA_RISK_COLORS:
                        _set_cell_shading(cell, CA_RISK_COLORS[risk].lstrip("#"))
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            run.font.bold = True
                elif risk in RISK_COLORS:
                    _set_cell_shading(cell, RISK_COLORS[risk].lstrip("#"))
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            run.font.bold = True


def _find_paragraph(doc: Document, text_fragment: str):
    """Return the first paragraph containing a stable template anchor."""
    return next((para for para in doc.paragraphs if text_fragment in para.text), None)


def _find_table(doc: Document, text_fragment: str):
    """Return the first template table containing a stable label."""
    for table in doc.tables:
        if any(text_fragment in cell.text for row in table.rows for cell in row.cells):
            return table
    return None


def _find_risk_summary_table(doc: Document):
    """Return the compact risk-count table shown in the Risk Classification section."""
    for table in doc.tables:
        if (
            len(table.rows) >= 8
            and len(table.columns) == 2
            and table.rows[0].cells[0].text.strip().lower() == "critical"
        ):
            return table
    return None


def _set_cell_no_borders(cell) -> None:
    """Remove the layout-only table borders without affecting its nested content."""
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "nil")
        tc_borders.append(border)
    cell._tc.get_or_add_tcPr().append(tc_borders)


def _copy_risk_cell_style(source, target) -> None:
    """Copy the colour treatment of a template risk-summary cell."""
    source_pr = source._tc.tcPr
    target_pr = target._tc.get_or_add_tcPr()
    for tag in ("w:shd", "w:tcBorders"):
        source_element = source_pr.find(qn(tag)) if source_pr is not None else None
        target_element = target_pr.find(qn(tag))
        if target_element is not None:
            target_pr.remove(target_element)
        if source_element is not None:
            target_pr.append(deepcopy(source_element))

    target.text = source.text
    source_run = source.paragraphs[0].runs[0] if source.paragraphs[0].runs else None
    for paragraph in target.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            run.font.name = "Cambria"
            run.font.size = Pt(10)
            if source_run is not None:
                run.font.bold = source_run.font.bold
                run.font.color.rgb = source_run.font.color.rgb


def _insert_chart_beside_risk_table(doc: Document, anchor_paragraph, table, image: io.BytesIO) -> None:
    """Place the Risk Classification chart and table below section 12's summary."""
    if image is None or table is None or anchor_paragraph is None:
        return

    # Use a single, compact table rather than nesting the source risk table in
    # another table. Word otherwise overestimates its height and moves it to a
    # new page even when there is visible room below the section narrative.
    layout = doc.add_table(rows=len(table.rows), cols=3)
    layout.autofit = False
    _set_table_width(layout, 20.0)
    _set_table_fixed_layout(layout)
    _set_column_widths(layout, [6600, 3000, 1740])
    for row in layout.rows:
        _prevent_row_split(row)
        _set_row_height(row, 18)

    chart_cell = layout.cell(0, 0).merge(layout.cell(len(table.rows) - 1, 0))
    _set_cell_no_borders(chart_cell)
    _set_cell_vertical_alignment(chart_cell)

    # The source table is defined earlier in the template, but this summary
    # belongs directly below the Vulnerabilities Details narrative. Insert the
    # compact replacement at that textual anchor and remove the original.
    anchor_paragraph._p.addnext(layout._tbl)
    for row_index, source_row in enumerate(table.rows):
        _copy_risk_cell_style(source_row.cells[0], layout.rows[row_index].cells[1])
        _copy_risk_cell_style(source_row.cells[1], layout.rows[row_index].cells[2])
        _set_cell_vertical_alignment(layout.rows[row_index].cells[1])
        _set_cell_vertical_alignment(layout.rows[row_index].cells[2])

    table._tbl.getparent().remove(table._tbl)

    chart_paragraph = chart_cell.paragraphs[0]
    chart_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # This fits below the section introduction on the same landscape page.
    chart_paragraph.add_run().add_picture(image, width=Inches(2.7), height=Inches(2.8))


def _insert_chart_after_table(doc: Document, table, image: io.BytesIO) -> None:
    """Insert the single executive-summary chart after its source table."""
    if image is None or table is None:
        return
    chart_paragraph = doc.add_paragraph()
    table._tbl.addnext(chart_paragraph._p)
    chart_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    chart_paragraph.add_run().add_picture(image, width=Inches(3.7), height=Inches(3.85))


def build_word_report(
    template_path: Path | str,
    output_path: Path | str,
    metadata: EngagementMetadata,
    va_df: pd.DataFrame,
    ca_df: pd.DataFrame,
    va_risk_summary: dict[str, int],
    ca_risk_summary: dict[str, int],
) -> Path:
    """Build the Word audit report.

    Parameters
    ----------
    template_path : Path
        Path to the Word template (.docx).
    output_path : Path
        Path where the output Word report will be saved.
    metadata : EngagementMetadata
        Engagement metadata.
    va_df : pd.DataFrame
        Processed VA data (normal report, with all columns).
    ca_df : pd.DataFrame
        Processed CA data (normal report, with all columns).
    va_risk_summary : dict
        VA risk level counts including "Grand Total".
    ca_risk_summary : dict
        CA risk level counts including "Grand Total".

    Returns
    -------
    Path
        Path to the saved Word report.
    """
    template_path = Path(template_path)
    output_path = Path(output_path)

    # Copy template to temp location to avoid locking
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp.close()
    shutil.copy2(template_path, tmp.name)
    doc = Document(tmp.name)

    # 1. Replace client name
    _replace_client_name(doc, metadata.client_name)

    # 1b. Populate document details (Released Date, etc.)
    _populate_document_details(doc, metadata)

    # 2. Update executive summary text
    _populate_executive_summary_text(doc, va_risk_summary, ca_risk_summary)

    # 3. Populate summary tables
    _populate_executive_summary_table(doc, va_risk_summary, ca_risk_summary)

    # 4. Add one pie chart after the executive-summary table.  The template's
    # normal page flow determines pagination; hard-coded page numbers produce
    # blank pages whenever the input lengths change.
    va_chart_buf = _generate_pie_chart_image(
        va_risk_summary, "Vulnerability Risk Distribution", RISK_COLORS, ca_risk_summary
    )
    _insert_chart_after_table(doc, _find_table(doc, "Security Assessment"), va_chart_buf)
    _insert_chart_beside_risk_table(
        doc,
        _find_paragraph(doc, "The audit reported"),
        _find_risk_summary_table(doc),
        va_chart_buf,
    )

    # 5. Find insertion points for VA and CA tables in VULNERABILITIES DETAILS
    va_anchor = _find_paragraph(doc, "The below table shows the detailed report of the VA scan done on assets.")
    ca_anchor = _find_paragraph(doc, "The below table shows the detailed report of the Compliance scan done on Assets.")

    # 6. Insert detailed tables at their own stable text anchors.  Holding the
    # paragraph objects avoids index drift after document elements are added.
    if va_anchor is not None and not va_df.empty:
        _create_detailed_table(
            doc, va_anchor, va_df, VA_COLUMNS,
            [4, 14, 23, 6, 10, 5, 20, 12, 8],
            {"Vulnerbility Title", "Description", "Recommendation ", "Reference"},
            top_left_columns={"Description", "Recommendation ", "Reference"},
        )
    if ca_anchor is not None and not ca_df.empty:
        _create_detailed_table(
            doc, ca_anchor, ca_df, CA_COLUMNS,
            [5, 20, 12, 31, 27, 8],
            {"Title", "Description", "Solution"},
            top_left_columns={"Description", "Solution"},
        )

    # 8. Save
    os.makedirs(output_path.parent, exist_ok=True)
    save_path = Path(output_path)
    for attempt in range(10):
        try:
            doc.save(str(save_path))
            break
        except PermissionError:
            if attempt < 9:
                stem = save_path.stem
                suffix = save_path.suffix
                # Strip existing numeric suffix (e.g. _2) before adding new one
                base = stem.rsplit("_", 1)[0] if stem.rsplit("_", 1)[-1].isdigit() else stem
                save_path = save_path.parent / f"{base}_{attempt + 2}{suffix}"
                logger.warning("File locked, retrying with: %s", save_path.name)
            else:
                raise

    # Cleanup temp file
    try:
        os.unlink(tmp.name)
    except OSError:
        pass

    logger.info("Word report saved: %s", save_path)
    return save_path
